from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

BRAND = RGBColor(0x4F, 0x46, 0xE5)
ACCENT = RGBColor(0x7C, 0x3A, 0xED)

def heading(text, level=1, color=BRAND):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(26)
    elif level == 1:
        run.font.size = Pt(17)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p

# Cover
heading('IdeaValidator', level=0)
tag = doc.add_paragraph()
r = tag.add_run('Launch, Marketing & Monetization Playbook')
r.italic = True
r.font.size = Pt(14)
r.font.color.rgb = ACCENT

para('A practical guide to going live, attracting your first 100 paying users, and growing IdeaValidator into a sustainable SaaS business.', italic=True)

doc.add_paragraph()

# Section 1
heading('1. Pre-Launch Checklist (Days 1-3)', level=1)
para('Before promoting your platform, complete this checklist to ensure a smooth user experience:')
for item in [
    'Deploy the live version of the app to your production URL.',
    'Test the full signup → idea submission → validation flow as a brand new user.',
    'Run a real Stripe checkout in test mode using card 4242 4242 4242 4242.',
    'Verify the Stripe webhook is receiving events (Stripe Dashboard → Webhooks → Events).',
    'Create 2-3 sample validation reports to use as demos and screenshots.',
    'Write a one-sentence elevator pitch you can repeat consistently.',
    'Prepare a 60-second screen-recorded demo video (use Loom or QuickTime).',
]:
    bullet(item)

# Section 2
heading('2. Your Positioning & Messaging', level=1)
para('Core Value Proposition', bold=True)
para('"Validate any business idea in 48 hours with AI-powered surveys, competitor research, and market sizing — for less than the cost of one bad decision."')

para('Target Audiences', bold=True)
for item in [
    'Aspiring entrepreneurs evaluating their first business idea.',
    'Side hustlers comparing multiple income stream ideas.',
    'Indie hackers and solo founders avoiding wasted build time.',
    'Small-business owners planning new products or services.',
]:
    bullet(item)

para('Key Messages That Convert', bold=True)
for item in [
    '"Stop guessing. Start validating." — speaks to fear of failure.',
    '"Skip months of research — get answers in 48 hours." — speaks to speed.',
    '"Free to start. Upgrade when you are serious." — removes friction.',
    '"AI-powered insights, founder-approved decisions." — credibility.',
]:
    bullet(item)

# Section 3
heading('3. Launch Week Marketing Plan (Days 4-10)', level=1)

heading('Day 1 — Product Hunt Launch', level=2)
for item in [
    'Submit at 12:01 AM PT on a Tuesday, Wednesday, or Thursday for max visibility.',
    'Prepare: tagline, 3 gallery images, 60-sec demo video, first comment from you.',
    'Notify 50+ contacts the day before to upvote in the first 4 hours.',
    'Reply to every comment personally within an hour.',
    'Goal: 300+ upvotes, top 5 of the day, 1,000+ visits.',
]:
    bullet(item)

heading('Day 2-3 — Indie Hacker Communities', level=2)
for item in [
    'Post a "Show IH" thread on indiehackers.com with your story and metrics.',
    'Share in r/SideProject, r/Entrepreneur, r/startups with genuine, non-spammy posts.',
    'Post in Hacker News Show HN (Tuesday-Thursday morning ET works best).',
    'Share in relevant Slack/Discord communities (Indie Worldwide, Makerlog).',
]:
    bullet(item)

heading('Day 4-5 — Social Media Blast', level=2)
for item in [
    'Twitter/X: Launch thread with your story, screenshots, and a clear CTA.',
    'LinkedIn: Founder story post targeting entrepreneurs and side hustlers.',
    'TikTok/Reels: Short demo videos showing the 48-hour validation in action.',
    'Tag relevant creators and ask them to try it (offer free Pro for feedback).',
]:
    bullet(item)

heading('Day 6-7 — Personal Network & Outreach', level=2)
for item in [
    'Email 50 friends, former colleagues, and mentors personally.',
    'Cold-email 25 entrepreneurship podcasters offering them a free Business plan.',
    'DM 10 micro-influencers (5K-50K followers) in the entrepreneurship space.',
]:
    bullet(item)

# Section 4
heading('4. Monetization Strategy', level=1)

para('Current Pricing Structure', bold=True)
tbl = doc.add_table(rows=4, cols=4)
tbl.style = 'Light Grid Accent 1'
hdr = tbl.rows[0].cells
for i, h in enumerate(['Plan', 'Price', 'Validations/Month', 'Target User']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
data = [
    ('Free', '$0', '2', 'Curious explorers, top-of-funnel'),
    ('Pro', '$29/mo', '15', 'Serious founders evaluating ideas'),
    ('Business', '$79/mo', 'Unlimited', 'Agencies, accelerators, power users'),
]
for i, row in enumerate(data, start=1):
    cells = tbl.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

doc.add_paragraph()

para('Conversion Optimization Tactics', bold=True)
for item in [
    'Show upgrade prompts the moment users hit the free limit (already implemented).',
    'Add a "What you would unlock" preview in the upgrade prompt.',
    'Send a 3-email onboarding sequence: welcome → first validation tips → upgrade offer.',
    'Offer 20% off the first month for new signups (limited-time popup).',
    'Add annual plans at 2 months free ($290/yr Pro, $790/yr Business) — boosts LTV.',
    'Display social proof: "X founders validated their ideas this week."',
]:
    bullet(item)

para('Future Revenue Streams', bold=True)
for item in [
    'One-time "Deep Dive" reports at $99 each for non-subscribers.',
    'White-label option for accelerators and incubators ($299/mo).',
    'Affiliate revenue from recommending tools (Stripe Atlas, Webflow, etc.).',
    'A founder community add-on ($19/mo) with peer feedback on validations.',
]:
    bullet(item)

# Section 5
heading('5. Growth Engine (Weeks 2-12)', level=1)

para('Content Marketing for SEO', bold=True)
for item in [
    'Publish 2 blog posts per week targeting long-tail keywords.',
    'Priority topics: "how to validate a business idea", "is my business idea good", "market research for startups".',
    'Create case studies of real validations (with permission) — these convert.',
    'Build a free resource library: validation checklists, survey templates, market sizing worksheets.',
]:
    bullet(item)

para('Community Building', bold=True)
for item in [
    'Launch a Discord or Slack for validated founders to share progress.',
    'Host a monthly live AMA with successful founders who validated with you.',
    'Create a public "Validated Ideas" gallery (opt-in, anonymized as needed).',
]:
    bullet(item)

para('Partnerships', bold=True)
for item in [
    'Partner with startup accelerators — offer free Business plans to their cohorts.',
    'Cross-promote with complementary tools (no-code builders, landing page tools).',
    'Guest post on top entrepreneurship blogs (StarterStory, Foundr, IndieHackers).',
    'Sponsor 2-3 niche entrepreneurship podcasts ($500-2K each — measure CAC).',
]:
    bullet(item)

para('Paid Acquisition (Once you have proof)', bold=True)
for item in [
    'Start with $500/mo Google Search Ads on high-intent keywords.',
    'Test Meta Ads with video demos targeting entrepreneurship interests.',
    'Track CAC (Customer Acquisition Cost) — keep it below 3x monthly subscription price.',
    'Scale spend only when LTV:CAC ratio exceeds 3:1.',
]:
    bullet(item)

# Section 6
heading('6. Metrics to Track from Day 1', level=1)
mt = doc.add_table(rows=7, cols=2)
mt.style = 'Light Grid Accent 1'
rows = [
    ('Metric', 'Why It Matters'),
    ('Signups per day', 'Top-of-funnel health.'),
    ('Free → Paid conversion %', 'Pricing/value alignment (target 3-5%).'),
    ('Time to first validation', 'Activation indicator — aim under 10 minutes.'),
    ('Monthly Recurring Revenue (MRR)', 'The North Star metric.'),
    ('Churn rate', 'Retention health (aim under 5%/mo).'),
    ('Customer Acquisition Cost (CAC)', 'Marketing efficiency.'),
]
for i, row in enumerate(rows):
    cells = mt.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val
        if i == 0:
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()

# Section 7
heading('7. 90-Day Milestone Roadmap', level=1)

heading('Month 1 — Launch & Validate', level=2)
for item in [
    'Goal: 500 signups, 15 paying customers, $500 MRR.',
    'Focus: Product Hunt, communities, founder outreach.',
    'Iterate based on first 50 user conversations.',
]:
    bullet(item)

heading('Month 2 — Optimize & Retain', level=2)
for item in [
    'Goal: 1,500 signups, 50 paying customers, $1,800 MRR.',
    'Launch email onboarding sequence and annual plans.',
    'Publish 8 SEO-optimized blog posts and 2 case studies.',
]:
    bullet(item)

heading('Month 3 — Scale Channels', level=2)
for item in [
    'Goal: 4,000 signups, 150 paying customers, $5,500 MRR.',
    'Start paid ads on validated keywords.',
    'Launch first accelerator partnership.',
    'Add annual plans, referral program, and affiliate program.',
]:
    bullet(item)

# Closing
heading('Final Reminders', level=1)
for item in [
    'Talk to your first 100 users personally — every conversation is gold.',
    'Ship improvements every week, not every quarter.',
    'Celebrate small wins publicly — momentum attracts users.',
    'Track everything, but obsess over the metrics that move revenue.',
]:
    bullet(item)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('You have built something real. Now go find the founders who need it.')
r.italic = True
r.bold = True
r.font.color.rgb = ACCENT
r.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = '/home/ubuntu/IdeaValidator_Launch_Marketing_Playbook.docx'
doc.save(out)
print('Saved:', out)
